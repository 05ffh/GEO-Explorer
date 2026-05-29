"""Generate actionable optimization plans as Markdown + PDF + DOCX reports."""

import os
import subprocess
import logging
from datetime import datetime
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import text

logger = logging.getLogger(__name__)


def _kpi_bar(value: float, width: int = 20) -> str:
    filled = int(abs(value) * width)
    return chr(9608) * filled + chr(9617) * (width - filled)


async def generate_optimization_plan(
    brand_name: str,
    collection_run_id: str,
    brand_id: str,
    db: AsyncSession,
    output_dir: str = "reports",
) -> dict:
    """Generate optimization plan in Markdown + PDF + DOCX formats."""
    # ── Fetch data ──
    rows = await db.execute(text("""
        SELECT cr.total_queries, cr.success_count, cr.failure_count,
               cr.collection_status, cr.analysis_status,
               ms.sov, ms.first_rec_rate, ms.accuracy_rate,
               ms.completeness_rate, ms.citation_rate, ms.sample_size, ms.details
        FROM collection_runs cr
        LEFT JOIN metrics_snapshots ms ON ms.collection_run_id = cr.id
        WHERE cr.id = :rid ORDER BY ms.created_at DESC LIMIT 1
    """), {"rid": collection_run_id})
    row = rows.fetchone()
    if not row:
        raise ValueError(f"Collection run {collection_run_id} not found")
    d = dict(row._mapping)

    ap_rows = await db.execute(text("""
        SELECT priority, action_type, suggested_content_type,
               ai_wrong_claims, correct_ground_truth, acceptance_criteria, status
        FROM action_plans WHERE brand_id = :bid AND status = 'pending'
        ORDER BY CASE priority WHEN 'P0' THEN 1 WHEN 'P1' THEN 2 ELSE 3 END LIMIT 30
    """), {"bid": brand_id})
    plans = [dict(r._mapping) for r in ap_rows.fetchall()]

    plat_rows = await db.execute(text("""
        SELECT platform, SUM(CASE WHEN status='success' THEN 1 ELSE 0 END) ok, COUNT(*) total
        FROM query_results WHERE collection_run_id=:rid GROUP BY platform ORDER BY platform
    """), {"rid": collection_run_id})
    platforms = [dict(r._mapping) for r in plat_rows.fetchall()]

    extended_kpis = (d.get("details") or {}).get("extended_kpis", {})

    hall_rows = await db.execute(text("""
        SELECT severity, COUNT(*) c FROM hallucination_results
        WHERE collection_run_id=:rid AND verdict='incorrect'
        GROUP BY severity ORDER BY severity
    """), {"rid": collection_run_id})
    hall_counts = {r.severity: r.c for r in hall_rows.fetchall()}

    from src.schemas.ground_truth import KPI_DISPLAY_NAMES

    # ── Build Markdown ──
    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    acc = d.get("accuracy_rate", 0) or 0
    sov = d.get("sov", 0) or 0
    frr = d.get("first_rec_rate", 0) or 0
    total_bad = sum(hall_counts.values())

    priority_groups = {"P0": [], "P1": [], "P2": []}
    for plan in plans:
        priority_groups[plan["priority"]].append(plan)

    lines = []
    lines.append(f"# {brand_name} AI 品牌认知优化方案")
    lines.append("")
    lines.append(f"**生成时间:** {now}  |  **采集:** {d['success_count']}/{d['total_queries']} 成功  |  **问题:** {total_bad} 条  |  **方案:** {len(plans)} 条")
    lines.append("")
    lines.append("## 一、执行摘要")
    lines.append(f"本次诊断发现 **{total_bad}** 条 AI 错误声明，声量 **{sov:.1%}**，准确率 **{acc:.1%}**。以下为可落地实施的优化方案。")
    lines.append("")

    lines.append("## 二、KPI 总览")
    lines.append("| 指标 | 得分 | 评估 |")
    lines.append("|------|------|------|")
    for key, val, note in [
        ("sov", sov, "需提升" if sov < 0.5 else "良好"),
        ("first_rec_rate", frr, "需场景化布局" if frr < 0.3 else "良好"),
        ("accuracy_rate", acc, "需纠正 AI 认知" if acc < 0.5 else "良好"),
        ("completeness_rate", d.get("completeness_rate", 0) or 0, ""),
        ("citation_rate", d.get("citation_rate", 0) or 0, ""),
    ]:
        name = KPI_DISPLAY_NAMES.get(key, key)
        lines.append(f"| {name} | {val:.1%} {_kpi_bar(val, 15)} | {note} |")
    for key, v in extended_kpis.items():
        name = KPI_DISPLAY_NAMES.get(key, key)
        val = v.get("value", 0)
        lines.append(f"| {name} | {val:.1%} {_kpi_bar(val, 15)} | |")
    lines.append("")

    lines.append("## 三、平台表现")
    lines.append("| 平台 | 成功率 |")
    lines.append("|------|--------|")
    for p in platforms:
        rate = p["ok"] / p["total"] * 100 if p["total"] else 0
        lines.append(f"| {p['platform']} | {rate:.0f}% ({p['ok']}/{p['total']}) |")
    lines.append("")

    lines.append(f"## 四、优化方案（共 {len(plans)} 条）")
    seq = 1
    for p_level, label, desc in [
        ("P0", "致命错误 — 24h 内修复", "以下问题可能导致 AI 给出严重失实的品牌描述。"),
        ("P1", "重要偏差 — 本周内修复", "以下问题影响品牌在 AI 中的认知质量。"),
        ("P2", "改善建议 — 本月内优化", "以下问题影响品牌丰富度表达。"),
    ]:
        group = priority_groups[p_level]
        if not group:
            continue
        lines.append(f"### {seq}. {label}")
        seq += 1
        lines.append(f"{desc}")
        for i, plan in enumerate(group[:10], 1):
            ai_claim = str(plan.get("ai_wrong_claims", {}).get("claim", ""))[:150]
            gt = str(plan.get("correct_ground_truth", {}).get("value", ""))[:150]
            field = plan.get("correct_ground_truth", {}).get("field", "")
            lines.append(f"**{seq-1}.{i} {plan['action_type']}: {field}**")
            lines.append(f"- AI 错误: {ai_claim}")
            lines.append(f"- 应修正为: {gt}")
            lines.append(f"- 内容类型: {plan.get('suggested_content_type', 'FAQ')}")
            lines.append(f"- 实施: 确认事实 → 更新GT → 生成FAQ/Schema → 发布 → 验证")
            lines.append("")

    lines.append("## 五、实施路线图")
    lines.append("| 阶段 | 时间 | 行动 | 预期 |")
    lines.append("|------|------|------|------|")
    lines.append(f"| 1 | 第1周 | 修复 {len(priority_groups['P0'])} 条 P0 | Accuracy → 50%+ |")
    lines.append(f"| 2 | 第2-3周 | 修复 {len(priority_groups['P1'])} 条 P1 | First Rec → 30%+ |")
    lines.append(f"| 3 | 第1-2月 | 优化 {len(priority_groups['P2'])} 条 P2 | SOV → 80%+ |")
    lines.append(f"| 持续 | 季度 | 定期诊断 | 维持品牌 AI 认知健康度 |")
    lines.append("")
    lines.append(f"*GEO Explorer Phase 10 | {now} | 数据来源: DeepSeek, Kimi, 豆包*")

    # ── Save Markdown ──
    os.makedirs(output_dir, exist_ok=True)
    date_str = datetime.now().strftime("%Y%m%d_%H%M")
    safe_name = brand_name.replace(" ", "_").replace("/", "_")
    md_path = os.path.join(output_dir, f"{safe_name}_{date_str}_优化方案.md")
    with open(md_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))

    # ── Generate PDF via md2pdf ──
    pdf_filename = f"{safe_name}_{date_str}_优化方案.pdf"
    pdf_path = os.path.join(output_dir, pdf_filename)
    try:
        home = os.path.expanduser("~")
        result = subprocess.run(
            ["md2pdf", md_path, pdf_path],
            capture_output=True, text=True, timeout=30,
            env={"HOME": home, "PATH": f"{home}/.nvm/versions/node/v22.22.2/bin:{os.environ.get('PATH', '')}",
                 "NVM_DIR": f"{home}/.nvm", **{k: v for k, v in os.environ.items() if k != 'PATH'}},
        )
        if result.returncode != 0:
            logger.error("md2pdf: %s", result.stderr[:200])
            pdf_path = None
    except Exception as e:
        logger.exception("md2pdf failed")
        pdf_path = None

    # ── Generate DOCX (editable Word) ──
    docx_path = os.path.join(output_dir, f"{safe_name}_{date_str}_优化方案.docx")
    try:
        _build_docx(brand_name, now, d, plans, priority_groups, platforms,
                    extended_kpis, hall_counts, KPI_DISPLAY_NAMES, docx_path)
    except Exception as e:
        logger.exception("docx generation failed")
        docx_path = None

    return {
        "markdown": os.path.abspath(md_path),
        "pdf": os.path.abspath(pdf_path) if pdf_path else None,
        "docx": os.path.abspath(docx_path) if docx_path else None,
        "action_count": len(plans),
        "p0_count": len(priority_groups["P0"]),
        "p1_count": len(priority_groups["P1"]),
    }


def _build_docx(brand_name, now, d, plans, priority_groups, platforms,
                extended_kpis, hall_counts, KPI_DISPLAY_NAMES, filepath):
    """Build an editable Word document with the optimization plan."""
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)

    # Title
    title = doc.add_heading(f'{brand_name} AI 品牌认知优化方案', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f'生成时间: {now}  |  采集状态: {d.get("collection_status", "N/A")}  '
                      f'|  {d["success_count"]}/{d["total_queries"]} 查询成功')
    doc.add_paragraph()

    # Executive Summary
    doc.add_heading('一、执行摘要', 1)
    acc = d.get("accuracy_rate", 0) or 0
    sov = d.get("sov", 0) or 0
    total_bad = sum(hall_counts.values())
    doc.add_paragraph(
        f'本次 GEO 诊断共覆盖 {len(platforms)} 个 AI 平台，发现 {total_bad} 条 AI 错误声明。'
        f'品牌声量份额 {sov:.1%}，准确率 {acc:.1%}。'
        f'以下为分优先级、可落地实施的优化执行方案。'
    )

    # KPI Table
    doc.add_heading('二、KPI 评分总览', 1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = '指标'
    hdr[1].text = '得分'
    hdr[2].text = '评估'
    kpi_data = [
        ("声量份额 (SOV)", sov, "良好" if sov > 0.5 else "需提升"),
        ("首次推荐率", d.get("first_rec_rate", 0) or 0, ""),
        ("准确率", acc, "需纠正 AI 认知" if acc < 0.5 else "良好"),
        ("完备性", d.get("completeness_rate", 0) or 0, ""),
        ("引用率", d.get("citation_rate", 0) or 0, ""),
    ]
    for key, val in list(extended_kpis.items())[:5]:
        name = KPI_DISPLAY_NAMES.get(key, key)
        kpi_data.append((name, val.get("value", 0), ""))
    for name, val, note in kpi_data:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = f'{val:.1%}'
        row[2].text = note

    # Platform
    doc.add_heading('三、平台采集表现', 1)
    pt = doc.add_table(rows=1, cols=2)
    pt.style = 'Light Grid Accent 1'
    pt.rows[0].cells[0].text = '平台'
    pt.rows[0].cells[1].text = '成功率'
    for p in platforms:
        row = pt.add_row().cells
        row[0].text = p['platform']
        rate = p['ok'] / p['total'] * 100 if p['total'] else 0
        row[1].text = f'{rate:.0f}% ({p["ok"]}/{p["total"]})'

    # Actions
    doc.add_heading(f'四、优化执行方案（共 {len(plans)} 条）', 1)
    p_labels = [("P0", "致命错误 — 24小时内修复"), ("P1", "重要偏差 — 本周内修复"), ("P2", "改善建议 — 本月内优化")]
    for p_level, section_title in p_labels:
        group = priority_groups[p_level]
        if not group:
            continue
        doc.add_heading(section_title, 2)
        for i, plan in enumerate(group[:10], 1):
            ai_claim = str(plan.get("ai_wrong_claims", {}).get("claim", ""))[:200]
            gt = str(plan.get("correct_ground_truth", {}).get("value", ""))[:200]
            field = plan.get("correct_ground_truth", {}).get("field", "")
            doc.add_heading(f'{p_level}.{i} {plan["action_type"]}: {field}', 3)
            p = doc.add_paragraph()
            p.add_run('AI 错误描述: ').bold = True
            p.add_run(ai_claim)
            p2 = doc.add_paragraph()
            p2.add_run('应修正为: ').bold = True
            p2.add_run(gt)
            p3 = doc.add_paragraph()
            p3.add_run('建议内容类型: ').bold = True
            p3.add_run(plan.get('suggested_content_type', 'FAQ'))
            doc.add_paragraph(
                '实施步骤: 1) 确认事实 → 2) 更新 GT → 3) 生成内容 → 4) 发布 → 5) 验证',
                style='List Bullet'
            )

    # Roadmap
    doc.add_heading('五、实施路线图', 1)
    rt = doc.add_table(rows=1, cols=4)
    rt.style = 'Light Grid Accent 1'
    hdr = rt.rows[0].cells
    hdr[0].text = '阶段'
    hdr[1].text = '时间'
    hdr[2].text = '行动'
    hdr[3].text = '预期效果'
    roadmap = [
        ("1", "第1周", f"修复 {len(priority_groups['P0'])} 条 P0", "Accuracy → 50%+"),
        ("2", "第2-3周", f"修复 {len(priority_groups['P1'])} 条 P1", "First Rec → 30%+"),
        ("3", "第1-2月", f"优化 {len(priority_groups['P2'])} 条 P2", "SOV → 80%+"),
        ("持续", "季度", "定期诊断", "维持 AI 认知健康度"),
    ]
    for stage, time, action, expected in roadmap:
        row = rt.add_row().cells
        row[0].text = stage
        row[1].text = time
        row[2].text = action
        row[3].text = expected

    doc.add_paragraph()
    doc.add_paragraph(f'GEO Explorer Phase 10 | {now} | 数据来源: DeepSeek, Kimi, 豆包, DuckDuckGo').italic = True

    doc.save(filepath)
