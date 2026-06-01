"""Async report generation Celery task (P1-6)."""
import asyncio
import hashlib
import logging
import os
import uuid
from datetime import datetime, timezone

from jinja2 import Environment, FileSystemLoader
from sqlalchemy.ext.asyncio import create_async_engine, async_sessionmaker
from sqlalchemy.pool import NullPool

from src.celery_app import app
from src.config import settings
from src.models.report_artifact import ReportArtifact
from src.reports.context_builder import build_report_context
from src.reports.sanitize_report import sanitize_report_context
from src.reports.report_quality import check_report_quality

logger = logging.getLogger(__name__)

engine = create_async_engine(settings.database_url, poolclass=NullPool)
SessionLocal = async_sessionmaker(engine, expire_on_commit=False)

TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), "templates")
_jinja_env = Environment(loader=FileSystemLoader(TEMPLATE_DIR))

REPORT_TEMPLATE_VERSION = "1.0"
CUSTOMER_LANGUAGE_VERSION = "1.0"


@app.task(bind=True, max_retries=2, acks_late=True, soft_time_limit=300, time_limit=420)
def report_generation_task(self, brand_id: str, org_id: str, collection_run_id: str,
                            editions: list | None = None, generated_by: str = "",
                            industry_template_id: str | None = None):
    """Async report generation — renders templates, exports PDF/DOCX, records artifacts."""
    if editions is None:
        editions = ["executive", "implementation", "customer"]

    async def _run():
        async with SessionLocal() as db:
            from sqlalchemy import select
            from src.models.brand import Brand

            brand_row = (await db.execute(
                select(Brand).where(Brand.id == uuid.UUID(brand_id))
            )).scalar_one_or_none()
            if not brand_row:
                return {"status": "error", "detail": "Brand not found"}

            brand = {"id": str(brand_row.id), "name": brand_row.name,
                     "industry": brand_row.industry or "", "website": getattr(brand_row, "website", "") or ""}

            industry_template = None
            if industry_template_id:
                from src.models.industry_template import IndustryTemplate
                industry_template = await db.get(IndustryTemplate, uuid.UUID(industry_template_id))

            results = {}
            for edition in editions:
                try:
                    artifact = await _generate_edition(
                        db, brand, collection_run_id, edition, generated_by,
                        industry_template,
                    )
                    results[edition] = artifact
                except Exception as exc:
                    logger.error(f"Report generation failed for {edition}: {exc}")
                    results[edition] = {"status": "failed", "error": str(exc)}

            return {"status": "completed", "results": results}

    return asyncio.run(_run())


async def _generate_edition(db, brand, collection_run_id, edition, generated_by, industry_template) -> dict:
    """Generate a single report edition and record the artifact."""
    now = datetime.now(timezone.utc)
    org_id = str(brand.get("org_id", ""))
    bid = str(brand["id"])

    # Build context
    ctx = await build_report_context(brand, collection_run_id, db, edition, industry_template)

    # Render template
    safe_name = brand["name"].replace(" ", "_").replace("/", "_")
    date_str = now.strftime("%Y%m%d_%H%M")

    for fmt in ["md", "pdf", "docx"]:
        gen_key = f"{bid}:{collection_run_id}:{edition}:{fmt}:{REPORT_TEMPLATE_VERSION}"

        # Idempotency: check existing
        existing = (await db.execute(
            __import__("sqlalchemy").select(ReportArtifact).where(
                ReportArtifact.generation_key == gen_key,
                ReportArtifact.status.in_(["generating", "queued"]),
            )
        )).scalar_one_or_none()
        if existing:
            continue

        # Create artifact record
        artifact = ReportArtifact(
            brand_id=uuid.UUID(bid),
            organization_id=uuid.UUID(org_id) if org_id else brand.get("organization_id", uuid.uuid4()),
            collection_run_id=uuid.UUID(collection_run_id),
            edition=edition,
            format=fmt,
            report_version=1,
            template_version=REPORT_TEMPLATE_VERSION,
            language_version=CUSTOMER_LANGUAGE_VERSION,
            industry_template_id=industry_template.id if industry_template else None,
            status="generating",
            generation_key=gen_key,
            generated_by=uuid.UUID(generated_by) if generated_by else None,
            locale="zh-CN",
        )
        db.add(artifact)
        await db.flush()

        try:
            out_dir = f"reports/{safe_name}_{date_str}"
            os.makedirs(out_dir, exist_ok=True)

            if fmt == "md":
                filepath = await _export_md(ctx, edition, out_dir, safe_name, date_str)
            elif fmt == "pdf":
                filepath = await _export_pdf(ctx, edition, out_dir, safe_name, date_str)
            elif fmt == "docx":
                filepath = await _export_docx(ctx, edition, out_dir, safe_name, date_str)
            else:
                continue

            if filepath and os.path.exists(filepath):
                artifact.file_path = filepath
                artifact.file_size_bytes = os.path.getsize(filepath)
                with open(filepath, "rb") as f:
                    artifact.file_hash = hashlib.sha256(f.read()).hexdigest()
                artifact.status = "generated"
                artifact.generated_at = now

                # Quality check
                if fmt == "md":
                    with open(filepath) as f:
                        qr = check_report_quality(f.read(), edition)
                    artifact.quality_status = qr.get("status", "warning")
                    artifact.quality_report_json = qr
                    if qr.get("status") == "failed":
                        artifact.status = "quality_failed"
            else:
                artifact.status = "failed"
                artifact.error_message = f"Export produced no file"

        except Exception as exc:
            artifact.status = "failed"
            artifact.error_message = str(exc)[:1000]
            logger.error(f"Report export failed for {edition}/{fmt}: {exc}")

        # Snapshot context (sanitized)
        artifact.context_snapshot = sanitize_report_context(ctx, edition)
        await db.commit()

        return {"artifact_id": str(artifact.id), "edition": edition, "format": fmt,
                "status": artifact.status, "path": artifact.file_path}

    return {"status": "generated", "edition": edition}


# ── Export helpers ───────────────────────────────────────────────────────────

async def _export_md(ctx, edition, out_dir, safe_name, date_str) -> str | None:
    template = _jinja_env.get_template(f"{edition}.md.j2")
    md_content = template.render(ctx=ctx)
    qr = check_report_quality(md_content, edition)
    if qr.get("status") == "failed":
        return None
    filename = f"{safe_name}_{date_str}_{edition}_v1.md"
    filepath = os.path.join(out_dir, filename)
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(md_content)
    return filepath


async def _export_pdf(ctx, edition, out_dir, safe_name, date_str) -> str | None:
    try:
        md_path = os.path.join(out_dir, f"{safe_name}_{date_str}_{edition}_v1.md")
        pdf_path = os.path.join(out_dir, f"{safe_name}_{date_str}_{edition}_v1.pdf")
        import subprocess
        home = os.path.expanduser("~")
        env = {"HOME": home, "PATH": f"{home}/.nvm/versions/node/v22.22.2/bin:{os.environ.get('PATH', '')}",
               "NVM_DIR": f"{home}/.nvm"}
        r = subprocess.run(["md2pdf", md_path, pdf_path], capture_output=True, text=True, timeout=30, env=env)
        if r.returncode != 0:
            logger.error(f"PDF export failed: {r.stderr}")
            return None
        return pdf_path if os.path.exists(pdf_path) else None
    except Exception:
        return None


async def _export_docx(ctx, edition, out_dir, safe_name, date_str) -> str | None:
    try:
        # Render as MD first, then convert
        template = _jinja_env.get_template(f"{edition}.md.j2")
        md_content = template.render(ctx=ctx)

        from docx import Document
        from docx.shared import Pt

        doc = Document()
        for line in md_content.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith("# "):
                doc.add_heading(stripped[2:], 0)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], 1)
            elif stripped.startswith("### "):
                doc.add_heading(stripped[4:], 2)
            elif stripped.startswith("#### "):
                doc.add_heading(stripped[5:], 3)
            else:
                p = doc.add_paragraph(stripped)
                p.style.font.size = Pt(10)

        filename = f"{safe_name}_{date_str}_{edition}_v1.docx"
        filepath = os.path.join(out_dir, filename)
        doc.save(filepath)
        return filepath
    except Exception:
        return None
