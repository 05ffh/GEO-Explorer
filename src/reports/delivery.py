"""Unified delivery orchestrator — produces a single folder with all reports for any brand."""

import os
import shutil
from datetime import datetime
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import text


async def deliver_all_reports(
    brand_name: str,
    brand_id: str,
    collection_run_id: str,
    db: AsyncSession,
    output_dir: str = "reports",
) -> dict:
    """Generate the complete report package for a brand diagnosis.

    Produces a single folder containing:
      README.md          — index with links and publish guide
      诊断报告.md         — KPI scores + hallucination summary
      优化方案.md/.docx/.pdf — 3-format optimization plan aligned with Content Packages
      NN_{theme}.md / NN_{theme}_schema.json — individual publishable content pieces
    """
    date_str = datetime.now().strftime("%Y%m%d_%H%M")
    safe_name = brand_name.replace(" ", "_").replace("/", "_")
    deliver_dir = os.path.join(output_dir, f"{safe_name}_{date_str}")
    os.makedirs(deliver_dir, exist_ok=True)

    # 1. Fetch analysis data
    analysis = await _fetch_analysis_data(brand_id, collection_run_id, brand_name, db)

    # 2. Diagnostic report (.md)
    diag_path = _write_diagnostic(brand_name, analysis, deliver_dir)

    # 3. Optimization plan (.md + .docx + .pdf) — aligned with Content Packages
    opt_result = await _write_optimization(brand_name, analysis, deliver_dir)

    # 4. Export individual Content Package files
    cp_files = await _export_content_pieces(brand_name, brand_id, deliver_dir, db)

    # 5. Build README index
    _write_readme(brand_name, analysis, deliver_dir, opt_result, cp_files)

    # ── Phase A: Format failure writeback ──
    import json as _json
    format_failures = []
    for fmt_label, path in [("md", diag_path), ("docx", opt_result.get("docx")), ("pdf", opt_result.get("pdf"))]:
        if not path or not os.path.exists(path):
            format_failures.append(fmt_label)
        elif os.path.getsize(path) == 0:
            format_failures.append(fmt_label)

    if format_failures:
        row = (await db.execute(text("""
            SELECT report_publishable, blocking_reasons_json, report_quality_summary_json
            FROM collection_runs WHERE id = :rid
        """), {"rid": collection_run_id})).fetchone()
        if row:
            rd = dict(row._mapping)
            br = rd.get("blocking_reasons_json") or []
            br.append({
                "code": "REPORT_FORMAT_FAILURE",
                "message": f"Report format(s) failed: {', '.join(format_failures)}",
                "severity": "block",
            })
            qs = rd.get("report_quality_summary_json") or {}
            qs["report_publishable"] = False
            qs["blocking_reasons"] = br
            await db.execute(text("""
                UPDATE collection_runs
                SET report_publishable = FALSE,
                    blocking_reasons_json = :br::jsonb,
                    report_quality_summary_json = :qs::jsonb
                WHERE id = :rid
            """), {"rid": collection_run_id, "br": _json.dumps(br), "qs": _json.dumps(qs)})
            await db.commit()

    return {
        "dir": os.path.abspath(deliver_dir),
        "diagnostic_md": diag_path,
        "optimization_md": opt_result.get("md"),
        "optimization_docx": opt_result.get("docx"),
        "optimization_pdf": opt_result.get("pdf"),
        "content_pieces": len(cp_files),
    }


# ── Internal helpers ──────────────────────────────────────────────────────────

async def _fetch_analysis_data(brand_id: str, collection_run_id: str, brand_name: str, db: AsyncSession) -> dict:
    """Gather all analysis data needed for reports."""
    data = {"brand_name": brand_name, "collection_run_id": collection_run_id}

    # Collection run summary
    cr = (await db.execute(text("""
        SELECT collection_status, analysis_status, total_queries, success_count,
               failure_count, created_at, collection_completed_at, report_quality_summary_json
        FROM collection_runs WHERE id = :rid
    """), {"rid": collection_run_id})).fetchone()
    if cr:
        data.update(dict(cr._mapping))

    # Ensure quality summary is always present
    if "report_quality_summary_json" not in data or data["report_quality_summary_json"] is None:
        data["report_quality_summary_json"] = {}

    # Metrics
    ms = (await db.execute(text("""
        SELECT sov, first_rec_rate, accuracy_rate, completeness_rate, citation_rate,
               sample_size, failure_rate, details
        FROM metrics_snapshots WHERE collection_run_id = :rid
        ORDER BY created_at DESC LIMIT 1
    """), {"rid": collection_run_id})).fetchone()
    if ms:
        data.update(dict(ms._mapping))

    # Per-platform
    plat = await db.execute(text("""
        SELECT platform, SUM(CASE WHEN status='success' THEN 1 ELSE 0 END) ok, COUNT(*) total
        FROM query_results WHERE collection_run_id=:rid GROUP BY platform ORDER BY platform
    """), {"rid": collection_run_id})
    data["platforms"] = [dict(r._mapping) for r in plat.fetchall()]

    # Hallucinations
    hall = await db.execute(text("""
        SELECT severity, verdict, COUNT(*) c FROM hallucination_results
        WHERE collection_run_id=:rid GROUP BY severity, verdict ORDER BY severity, verdict
    """), {"rid": collection_run_id})
    data["hallucinations"] = [dict(r._mapping) for r in hall.fetchall()]
    data["hall_total"] = sum(h["c"] for h in data["hallucinations"])
    data["hall_incorrect"] = sum(h["c"] for h in data["hallucinations"] if h["verdict"] == "incorrect")

    # Action plans
    ap = await db.execute(text("""
        SELECT priority, COUNT(*) c FROM action_plans
        WHERE brand_id=:bid GROUP BY priority ORDER BY priority
    """), {"bid": brand_id})
    data["action_plans"] = [dict(r._mapping) for r in ap.fetchall()]
    data["ap_total"] = sum(a["c"] for a in data["action_plans"])

    # Content Packages (already generated)
    cp = await db.execute(text("""
        SELECT cp.content_items, cp.schema_items, cp.fact_check_report,
               cp.publishing_checklist, cp.status
        FROM content_packages cp WHERE cp.brand_id = :bid
        ORDER BY cp.created_at
    """), {"bid": brand_id})
    data["content_packages"] = [dict(r._mapping) for r in cp.fetchall()]

    return data


def _render_quality_summary_section(analysis: dict) -> str:
    """Render ReportQualitySummary as Markdown for report front page."""
    qs = analysis.get("report_quality_summary_json", {})
    if not qs:
        return ""
    ai = qs.get("ai_hallucination", {})
    tmpl = qs.get("template_issue", {})
    gt = qs.get("gt_insufficient", {})
    nab = qs.get("not_about_brand", {})
    pub = qs.get("report_publishable", False)
    blocking = qs.get("blocking_reasons", [])

    lines = [
        "## 本次诊断可信度概览",
        "",
        "| 类别 | 数量 | 说明 |",
        "|------|:--:|------|",
        f"| AI 幻觉 (P0) | {ai.get('p0_count', 0)} | {ai.get('p0_explanation', '品牌核心事实与 GT 矛盾')} |",
        f"| AI 幻觉 (P1/P2) | {ai.get('p1_count', 0)}/{ai.get('p2_count', 0)} | 次要/边缘事实偏差 |",
        f"| 模板问题 | {tmpl.get('invalid_template_count', 0)} | 模板变量未替换，不计入 AI 幻觉 |",
        f"| GT 不足 | {gt.get('unsupported_claim_count', 0)} | GT 数据不足以核验，不计入 AI 幻觉 |",
        f"| 回答无关 | {nab.get('generic_statement_count', 0)} | 回答未涉及目标品牌，不计入 AI 幻觉 |",
        "",
        f"**报告状态:** {'可发布' if pub else '不可发布'}",
        "",
    ]
    if blocking:
        lines.append("**阻断/警告详情:**")
        for b in blocking:
            icon = "[BLOCK]" if b.get("severity") == "block" else "[WARN]"
            lines.append(f"- {icon} {b.get('code')}: {b.get('message')}")
        lines.append("")
    lines.append(f"> {ai.get('excluded_explanation', '模板问题、GT不足、回答无关不计入 AI 幻觉')}")
    return "\n".join(lines)


def _write_diagnostic(brand_name: str, data: dict, deliver_dir: str) -> str:
    """Write 诊断报告.md."""
    from src.schemas.ground_truth import KPI_DISPLAY_NAMES
    now = datetime.now().strftime("%Y-%m-%d %H:%M")

    def bar(v, w=20): return "█" * int(abs(v) * w) + "░" * (w - int(abs(v) * w))

    lines = [
        f"# {brand_name} GEO 诊断报告",
        f"",
        f"**生成时间:** {now}  ",
        f"**采集:** {data.get('collection_status','?')} | {data.get('success_count',0)}/{data.get('total_queries',0)} 成功 | {data.get('failure_count',0)} 失败",
        f"",
    ]

    # Prepend quality summary if available
    quality_section = _render_quality_summary_section(data)
    if quality_section:
        lines.append(quality_section)
        lines.append("")

    lines.append("## KPI 评分")
    lines.append("| 指标 | 得分 | |")
    lines.append("|------|------|--|")
    kpis = [
        ("sov", data.get("sov") or 0), ("first_rec_rate", data.get("first_rec_rate") or 0),
        ("accuracy_rate", data.get("accuracy_rate") or 0), ("completeness_rate", data.get("completeness_rate") or 0),
        ("citation_rate", data.get("citation_rate") or 0),
    ]
    for k, v in kpis:
        lines.append(f"| {KPI_DISPLAY_NAMES.get(k,k)} | {v:.1%} | {bar(v,15)} |")
    ek = (data.get("details") or {}).get("extended_kpis", {})
    for k, v in ek.items():
        lines.append(f"| {KPI_DISPLAY_NAMES.get(k,k)} | {v.get('value',0):.1%} | {bar(v.get('value',0),15)} |")
    lines.append("")

    # Platform
    lines.append("## 平台采集")
    lines.append("| 平台 | 成功 | 总数 |")
    lines.append("|------|------|------|")
    for p in data.get("platforms", []):
        lines.append(f"| {p['platform']} | {p['ok']} | {p['total']} |")
    lines.append("")

    # Hallucinations
    lines.append(f"## 幻觉检测: {data.get('hall_total',0)} 条")
    if data.get("hallucinations"):
        lines.append("| 严重度 | 判定 | 数量 |")
        lines.append("|--------|------|------|")
        sev = {"P0":"致命","P1":"重要","P2":"改善"}
        vl = {"correct":"✓","incorrect":"✗","uncertain":"?"}
        for h in data["hallucinations"]:
            lines.append(f"| {sev.get(h['severity'],h['severity'])} | {vl.get(h['verdict'],h['verdict'])} | {h['c']} |")
    lines.append("")

    lines.append(f"## Action Plans: {data.get('ap_total',0)} 条")
    lines.append(f"## Content Packages: {len(data.get('content_packages',[]))} 个主题")
    lines.append("")

    path = os.path.join(deliver_dir, "诊断报告.md")
    with open(path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    return path


async def _write_optimization(brand_name: str, data: dict, deliver_dir: str) -> dict:
    """Write 优化方案 in .md + .docx + .pdf, content aligned with Content Packages."""
    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    cp_list = data.get("content_packages", [])
    acc = data.get("accuracy_rate") or 0
    sov = data.get("sov") or 0

    # ── Markdown ──
    md_lines = [
        f"# {brand_name} AI 品牌认知优化方案",
        f"",
        f"**生成时间:** {now} | **KPI:** 准确率 {acc:.1%} | 声量 {sov:.1%} | 幻觉 {data.get('hall_incorrect',0)} 条",
        f"",
        f"## 执行摘要",
        f"本次 GEO 诊断共覆盖 {len(data.get('platforms',[]))} 个 AI 平台，发现 {data.get('hall_incorrect',0)} 条与事实不符的 AI 声明。",
        f"以下 {len(cp_list)} 个内容主题基于 GT 生成，可直接发布到官网以纠正 AI 认知。",
        f"",
    ]

    for i, cp in enumerate(cp_list):
        item = cp["content_items"][0]
        theme = item.get("theme", f"主题 {i+1}")
        body = item.get("body", "")
        if isinstance(body, list) and body:
            body = body[0].get("body", "")
        elif isinstance(body, dict):
            body = body.get("body", "")
        fields = item.get("source_fields", [])

        md_lines.append(f"## {i+1}. {theme}")
        md_lines.append(f"")
        md_lines.append(f"**覆盖 GT 字段:** {', '.join(fields)}")
        md_lines.append(f"")
        md_lines.append(str(body))
        md_lines.append(f"")

        # Publishing checklist
        checklist = cp.get("publishing_checklist", [])
        if checklist:
            md_lines.append(f"### 发布检查清单")
            for check in checklist:
                checked = "✓" if check.get("checked") else "☐"
                md_lines.append(f"- [{checked}] {check.get('item')}")
        md_lines.append("")

    md_path = os.path.join(deliver_dir, "优化方案.md")
    with open(md_path, "w", encoding="utf-8") as f:
        f.write("\n".join(md_lines))

    # ── DOCX ──
    docx_path = os.path.join(deliver_dir, "优化方案.docx")
    try:
        _build_optimization_docx(brand_name, now, data, cp_list, docx_path)
    except Exception:
        docx_path = None

    # ── PDF ──
    pdf_path = os.path.join(deliver_dir, "优化方案.pdf")
    try:
        import subprocess
        home = os.path.expanduser("~")
        r = subprocess.run(["md2pdf", md_path, pdf_path], capture_output=True, text=True, timeout=30,
                           env={"HOME": home, "PATH": f"{home}/.nvm/versions/node/v22.22.2/bin:{os.environ.get('PATH','')}",
                                "NVM_DIR": f"{home}/.nvm"})
        if r.returncode != 0:
            pdf_path = None
    except Exception:
        pdf_path = None

    return {"md": md_path, "docx": docx_path, "pdf": pdf_path}


def _build_optimization_docx(brand_name: str, now: str, data: dict, cp_list: list, filepath: str):
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles['Normal']
    style.font.size = Pt(10)

    title = doc.add_heading(f'{brand_name} AI 品牌认知优化方案', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'生成时间: {now} | 准确率: {(data.get("accuracy_rate") or 0):.1%} | 声量: {(data.get("sov") or 0):.1%}')

    doc.add_heading('执行摘要', 1)
    doc.add_paragraph(f'本次诊断共覆盖 {len(data.get("platforms",[]))} 个 AI 平台，发现 {data.get("hall_incorrect",0)} 条错误声明。'
                      f'以下 {len(cp_list)} 个内容主题可直接发布以纠正 AI 认知。')

    for i, cp in enumerate(cp_list):
        item = cp["content_items"][0]
        theme = item.get("theme", f"主题 {i+1}")
        body = item.get("body", "")
        if isinstance(body, list) and body:
            body = body[0].get("body", "")
        elif isinstance(body, dict):
            body = body.get("body", "")
        fields = item.get("source_fields", [])

        doc.add_heading(f'{i+1}. {theme}', 1)
        doc.add_paragraph(f'覆盖 GT 字段: {", ".join(fields)}').italic = True

        for paragraph in str(body).split('\n'):
            p = paragraph.strip()
            if not p:
                continue
            if p.startswith('#'):
                level = min(len(p) - len(p.lstrip('#')), 3)
                doc.add_heading(p.lstrip('#').strip(), level)
            else:
                doc.add_paragraph(p)

        checklist = cp.get("publishing_checklist", [])
        if checklist:
            doc.add_heading('发布检查清单', 3)
            for check in checklist:
                checked = "✓" if check.get("checked") else "☐"
                doc.add_paragraph(f'{checked} {check.get("item")}', style='List Bullet')

    doc.save(filepath)


async def _export_content_pieces(brand_name: str, brand_id: str, deliver_dir: str, db: AsyncSession) -> list:
    """Export individual .md + .json files per Content Package theme."""
    import json

    pkgs = (await db.execute(text("""
        SELECT content_items, schema_items, publishing_checklist
        FROM content_packages WHERE brand_id = :bid ORDER BY created_at
    """), {"bid": brand_id})).fetchall()

    files = []
    for i, pkg in enumerate(pkgs):
        d = dict(pkg._mapping)
        item = d["content_items"][0]
        theme = item.get("theme", f"package_{i+1}")
        safe = theme.replace(" ", "_").replace("(", "").replace(")", "").replace("&", "and")

        body = item.get("body", "")
        if isinstance(body, list) and body:
            body = body[0].get("body", "")
        elif isinstance(body, dict):
            body = body.get("body", "")

        # .md
        md_name = f"{i+1:02d}_{safe}.md"
        md_path = os.path.join(deliver_dir, md_name)
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(f"# {brand_name} - {theme}\n\n{body}\n")

        # .json
        json_name = f"{i+1:02d}_{safe}_schema.json"
        json_path = os.path.join(deliver_dir, json_name)
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(d["schema_items"], f, ensure_ascii=False, indent=2)

        files.append({"theme": theme, "md": md_name, "json": json_name, "len": len(body)})

    return files


def _write_readme(brand_name: str, data: dict, deliver_dir: str, opt_result: dict, cp_files: list):
    """Write README.md index."""
    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    lines = [
        f"# {brand_name} GEO 诊断交付物",
        f"",
        f"**生成时间:** {now} | **GEO Explorer Phase 10**",
        f"",
        f"## 文件说明",
        f"",
        f"| 文件 | 格式 | 用途 |",
        f"|------|------|------|",
        f"| 诊断报告.md | Markdown | KPI 评分 + 检测摘要 |",
        f"| 优化方案.md | Markdown | 可编辑优化方案 |",
        f"| 优化方案.docx | Word | 企业交付（可编辑） |",
        f"| 优化方案.pdf | PDF | 企业交付（保持排版） |",
    ]
    for f in cp_files:
        lines.append(f"| {f['md']} | Markdown | 可发布内容 - {f['theme']} |")
        lines.append(f"| {f['json']} | JSON-LD | Schema.org 结构化数据 - {f['theme']} |")

    lines.extend([
        "",
        "## 数据来源",
        f"- 采集查询: {data.get('success_count',0)}/{data.get('total_queries',0)} 成功",
        f"- AI 平台: {', '.join(p['platform'] for p in data.get('platforms',[]) if p['ok'] > 0)}",
        f"- 幻觉检测: {data.get('hall_total',0)} 条声明, {data.get('hall_incorrect',0)} 条错误",
        f"- Action Plans: {data.get('ap_total',0)} 条",
        f"- Content Packages: {len(cp_files)} 个主题",
        "",
        "## 发布流程",
        "1. 审阅 `诊断报告.md` 了解整体情况",
        "2. 逐主题审核 `优化方案.md` 或 `.docx` 中的内容",
        "3. 将对应的 `.md` 内容发布到官网",
        "4. 将 `.json` Schema.org 数据嵌入页面 `<head>`",
        "5. 验证 → 提交 URL → 等待 AI 认知改善",
        f"",
        f"*GEO Explorer Phase 10 | {now}*",
    ])

    with open(os.path.join(deliver_dir, "README.md"), "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
